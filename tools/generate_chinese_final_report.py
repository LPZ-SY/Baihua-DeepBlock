from __future__ import annotations

import csv
import json
from pathlib import Path
from typing import Iterable, Sequence

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
RESULT_ROOT = ROOT / "results" / "experiments" / "qrf_formal_hardware_matrix_v2"
OUTPUT_DIR = ROOT / "docs" / "Quantum_Route_Forge_正式实验最终报告_图表"
DOCX_PATH = ROOT / "docs" / "Quantum_Route_Forge_正式实验最终报告_中文版.docx"

FONT_REGULAR = Path(r"C:\Windows\Fonts\msyh.ttc")
FONT_BOLD = Path(r"C:\Windows\Fonts\msyhbd.ttc")

NAVY = "17365D"
BLUE = "2E74B5"
LIGHT_BLUE = "DCE6F1"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "667085"
DARK = "1F2937"
GREEN = "2E7D5B"
GOLD = "B7791F"
RED = "B42318"
WHITE = "FFFFFF"

PIL_COLORS = {
    "blue": (46, 116, 181),
    "navy": (23, 54, 93),
    "orange": (230, 126, 34),
    "green": (46, 125, 91),
    "gold": (183, 121, 31),
    "red": (180, 35, 24),
    "gray": (102, 112, 133),
    "light_gray": (234, 236, 240),
    "dark": (31, 41, 55),
}


def load_json(path: Path):
    return json.loads(path.read_text(encoding="utf-8"))


def load_csv(path: Path) -> list[dict[str, str]]:
    with path.open("r", encoding="utf-8-sig", newline="") as handle:
        return list(csv.DictReader(handle))


def pil_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    return ImageFont.truetype(str(FONT_BOLD if bold else FONT_REGULAR), size=size)


def centered_text(draw: ImageDraw.ImageDraw, xy: tuple[int, int], text: str, font, fill) -> None:
    draw.text(xy, text, font=font, fill=fill, anchor="mm")


def create_canvas(title: str, subtitle: str | None = None):
    image = Image.new("RGB", (2400, 1350), "white")
    draw = ImageDraw.Draw(image)
    centered_text(draw, (1200, 62), title, pil_font(58, True), PIL_COLORS["navy"])
    if subtitle:
        centered_text(draw, (1200, 125), subtitle, pil_font(29), PIL_COLORS["gray"])
    return image, draw


def draw_y_axis(
    draw: ImageDraw.ImageDraw,
    left: int,
    top: int,
    right: int,
    bottom: int,
    maximum: float,
    ticks: Sequence[float],
    label: str,
) -> None:
    axis_font = pil_font(28)
    draw.line((left, top, left, bottom), fill=PIL_COLORS["dark"], width=3)
    draw.line((left, bottom, right, bottom), fill=PIL_COLORS["dark"], width=3)
    for tick in ticks:
        y = bottom - int((tick / maximum) * (bottom - top))
        draw.line((left, y, right, y), fill=PIL_COLORS["light_gray"], width=2)
        draw.text((left - 24, y), f"{tick:.2f}", font=axis_font, fill=PIL_COLORS["gray"], anchor="rm")
    draw.text((95, (top + bottom) // 2), label, font=pil_font(30, True), fill=PIL_COLORS["dark"], anchor="mm")


def save_chart(image: Image.Image, filename: str) -> Path:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    path = OUTPUT_DIR / filename
    image.save(path, format="PNG", dpi=(300, 300), optimize=True)
    return path


def chart_task_quality(rows: list[dict[str, str]]) -> Path:
    image, draw = create_canvas(
        "图1  24个真机任务的质量命中率",
        "实测量子分布与冻结均匀随机参考的逐任务比较",
    )
    left, top, right, bottom = 190, 205, 2320, 1080
    maximum = 0.42
    draw_y_axis(draw, left, top, right, bottom, maximum, [0, 0.1, 0.2, 0.3, 0.4], "命中率")
    group_w = (right - left) / len(rows)
    bar_w = max(20, int(group_w * 0.31))
    label_font = pil_font(23)
    for idx, row in enumerate(rows, start=1):
        center = left + (idx - 0.5) * group_w
        quantum = float(row["quality_hit_rate"])
        random = float(row["random_quality_hit_rate"])
        q_y = bottom - int(quantum / maximum * (bottom - top))
        r_y = bottom - int(random / maximum * (bottom - top))
        draw.rectangle((center - bar_w - 2, q_y, center - 2, bottom), fill=PIL_COLORS["blue"])
        draw.rectangle((center + 2, r_y, center + bar_w + 2, bottom), fill=PIL_COLORS["orange"])
        draw.text((center, bottom + 18), str(idx), font=label_font, fill=PIL_COLORS["dark"], anchor="ma")
    draw.text((1200, 1155), "任务序号（完整任务ID见报告附录）", font=pil_font(28), fill=PIL_COLORS["dark"], anchor="mm")
    draw.rectangle((730, 1199, 770, 1239), fill=PIL_COLORS["blue"])
    draw.text((785, 1219), "真机QHR", font=pil_font(27), fill=PIL_COLORS["dark"], anchor="lm")
    draw.rectangle((1130, 1199, 1170, 1239), fill=PIL_COLORS["orange"])
    draw.text((1185, 1219), "随机参考", font=pil_font(27), fill=PIL_COLORS["dark"], anchor="lm")
    centered_text(draw, (1200, 1300), "结果：24个任务中，真机QHR超过随机参考的任务数为0。", pil_font(28, True), PIL_COLORS["red"])
    return save_chart(image, "图1_逐任务量子与随机命中率.png")


def draw_error_bar(draw, x: int, mean: float, ci: Sequence[float], maximum: float, top: int, bottom: int, color) -> None:
    def y_of(value: float) -> int:
        return bottom - int(value / maximum * (bottom - top))

    y_low, y_high = y_of(float(ci[0])), y_of(float(ci[1]))
    draw.line((x, y_low, x, y_high), fill=color, width=5)
    draw.line((x - 14, y_low, x + 14, y_low), fill=color, width=5)
    draw.line((x - 14, y_high, x + 14, y_high), fill=color, width=5)


def chart_backend_summary(stats: dict) -> Path:
    image, draw = create_canvas(
        "图2  按后端分层的质量命中率",
        "误差线为以硬件任务为重复单位的Bootstrap 95%置信区间；不用于后端排名",
    )
    left, top, right, bottom = 220, 220, 2250, 1050
    maximum = 0.36
    draw_y_axis(draw, left, top, right, bottom, maximum, [0, 0.1, 0.2, 0.3], "平均命中率")
    backends = ["Baihua", "Dongling", "Shenglian"]
    centers = [560, 1230, 1900]
    bar_w = 120
    for backend, center in zip(backends, centers):
        entry = stats["by_backend"][backend]
        q = entry["quality_hit_rate"]
        r = entry["random_quality_hit_rate"]
        q_y = bottom - int(float(q["mean"]) / maximum * (bottom - top))
        r_y = bottom - int(float(r["mean"]) / maximum * (bottom - top))
        draw.rectangle((center - 150, q_y, center - 30, bottom), fill=PIL_COLORS["blue"])
        draw.rectangle((center + 30, r_y, center + 150, bottom), fill=PIL_COLORS["orange"])
        draw_error_bar(draw, center - 90, q["mean"], q["bootstrap_95_ci"], maximum, top, bottom, PIL_COLORS["navy"])
        draw_error_bar(draw, center + 90, r["mean"], r["bootstrap_95_ci"], maximum, top, bottom, PIL_COLORS["gold"])
        draw.text((center - 90, q_y - 35), f"{q['mean']:.3f}", font=pil_font(26, True), fill=PIL_COLORS["blue"], anchor="ms")
        draw.text((center + 90, r_y - 35), f"{r['mean']:.3f}", font=pil_font(26, True), fill=PIL_COLORS["orange"], anchor="ms")
        draw.text((center, bottom + 42), backend, font=pil_font(34, True), fill=PIL_COLORS["dark"], anchor="ma")
    draw.rectangle((790, 1150, 840, 1200), fill=PIL_COLORS["blue"])
    draw.text((860, 1175), "真机QHR", font=pil_font(29), fill=PIL_COLORS["dark"], anchor="lm")
    draw.rectangle((1190, 1150, 1240, 1200), fill=PIL_COLORS["orange"])
    draw.text((1260, 1175), "随机参考", font=pil_font(29), fill=PIL_COLORS["dark"], anchor="lm")
    centered_text(draw, (1200, 1290), "三个后端的描述性均值均低于对应随机参考。", pil_font(29, True), PIL_COLORS["red"])
    return save_chart(image, "图2_按后端质量命中率.png")


def chart_bqm_outcomes(rows: list[dict[str, str]], hybrid_rows: list[dict[str, str]]) -> Path:
    exact_hits = sum(float(row["best_gap"]) == 0.0 for row in rows)
    strict_wins = sum(float(row["strict_improvement_rate"]) > 0.0 for row in rows)
    source_wins = sum(str(row["quantum_source_win"]).lower() == "true" for row in hybrid_rows)
    route_wins = sum(float(row["delta_QR"]) > 0.0 for row in hybrid_rows)
    labels = [
        "命中全局最优BQM能量",
        "严格优于同预算经典阈值",
        "量子来源进入C+Q最终选择",
        "相对C+R改善最终路线距离",
    ]
    values = [exact_hits, strict_wins, source_wins, route_wins]
    colors = [PIL_COLORS["green"], PIL_COLORS["red"], PIL_COLORS["blue"], PIL_COLORS["red"]]
    image, draw = create_canvas(
        "图3  真机BQM能量与最终路线结果总览",
        "计数单位为硬件任务；每项总任务数均为24",
    )
    left, right = 850, 2200
    top, row_h = 265, 205
    axis_y = top + row_h * 4 + 25
    draw.line((left, axis_y, right, axis_y), fill=PIL_COLORS["dark"], width=3)
    for tick in [0, 6, 12, 18, 24]:
        x = left + int((right - left) * tick / 24)
        draw.line((x, top - 40, x, axis_y), fill=PIL_COLORS["light_gray"], width=2)
        draw.text((x, axis_y + 24), str(tick), font=pil_font(28), fill=PIL_COLORS["gray"], anchor="ma")
    for idx, (label, value, color) in enumerate(zip(labels, values, colors)):
        y = top + idx * row_h
        draw.text((790, y + 45), label, font=pil_font(34, True), fill=PIL_COLORS["dark"], anchor="rm")
        end_x = left + int((right - left) * value / 24)
        if value > 0:
            draw.rounded_rectangle((left, y, end_x, y + 90), radius=20, fill=color)
        else:
            draw.ellipse((left - 9, y + 35, left + 9, y + 53), fill=color)
        draw.text((max(end_x + 28, left + 28), y + 45), f"{value}/24", font=pil_font(36, True), fill=color, anchor="lm")
    centered_text(
        draw,
        (1200, 1260),
        "结论：真机能命中最优BQM能量，但经典同预算阈值已达到同一能量，因此没有严格能量胜出。",
        pil_font(28, True),
        PIL_COLORS["navy"],
    )
    return save_chart(image, "图3_BQM能量结果总览.png")


def chart_classical_threshold(stats: dict) -> Path:
    image, draw = create_canvas(
        "图4  按实例统计的经典阈值达到率与严格改进率",
        "达到阈值允许能量相等；严格改进要求能量显著低于同预算经典阈值",
    )
    left, top, right, bottom = 220, 220, 2250, 1050
    maximum = 0.09
    draw_y_axis(draw, left, top, right, bottom, maximum, [0, 0.02, 0.04, 0.06, 0.08], "平均比率")
    instances = [
        ("seed2026_c4_v2_medium", "2026-C4\n中等压力"),
        ("seed2027_c4_v2_tight", "2027-C4\n紧张压力"),
        ("seed2026_c6_v2_medium", "2026-C6\n中等压力"),
        ("seed2027_c6_v2_tight", "2027-C6\n紧张压力"),
    ]
    centers = [460, 970, 1480, 1990]
    for (key, label), center in zip(instances, centers):
        entry = stats["by_instance"][key]
        reach = float(entry["classical_reach_feasible_rate"]["mean"])
        strict = float(entry["strict_improvement_rate"]["mean"])
        reach_y = bottom - int(reach / maximum * (bottom - top))
        strict_y = bottom - int(strict / maximum * (bottom - top))
        draw.rectangle((center - 130, reach_y, center - 20, bottom), fill=PIL_COLORS["green"])
        if strict > 0:
            draw.rectangle((center + 20, strict_y, center + 130, bottom), fill=PIL_COLORS["red"])
        else:
            draw.ellipse((center + 66, bottom - 9, center + 84, bottom + 9), fill=PIL_COLORS["red"])
        draw.text((center - 75, reach_y - 28), f"{reach:.3f}", font=pil_font(26, True), fill=PIL_COLORS["green"], anchor="ms")
        draw.text((center + 75, bottom - 28), "0", font=pil_font(26, True), fill=PIL_COLORS["red"], anchor="ms")
        draw.multiline_text((center, bottom + 45), label, font=pil_font(29, True), fill=PIL_COLORS["dark"], anchor="ma", align="center", spacing=8)
    draw.rectangle((780, 1185, 830, 1235), fill=PIL_COLORS["green"])
    draw.text((850, 1210), "达到经典阈值", font=pil_font(28), fill=PIL_COLORS["dark"], anchor="lm")
    draw.ellipse((1245, 1201, 1263, 1219), fill=PIL_COLORS["red"])
    draw.text((1280, 1210), "严格优于经典阈值", font=pil_font(28), fill=PIL_COLORS["dark"], anchor="lm")
    centered_text(draw, (1200, 1305), "所有实例均有真机样本达到阈值，但严格改进率均为0。", pil_font(28, True), PIL_COLORS["navy"])
    return save_chart(image, "图4_经典阈值达到率与严格改进率.png")


def chart_hybrid_delta(hybrid_rows: list[dict[str, str]]) -> Path:
    image, draw = create_canvas(
        "图5  C+Q相对C+R和纯经典的路线距离增量",
        "正值代表加入真机量子候选后路线更短；统计单位为硬件任务",
    )
    left, top, right, bottom = 190, 220, 2320, 1030
    minimum, maximum = -0.02, 0.02
    draw.line((left, top, left, bottom), fill=PIL_COLORS["dark"], width=3)
    draw.line((left, bottom, right, bottom), fill=PIL_COLORS["dark"], width=3)
    for tick in [-0.02, -0.01, 0, 0.01, 0.02]:
        y = bottom - int((tick - minimum) / (maximum - minimum) * (bottom - top))
        width = 5 if tick == 0 else 2
        color = PIL_COLORS["dark"] if tick == 0 else PIL_COLORS["light_gray"]
        draw.line((left, y, right, y), fill=color, width=width)
        draw.text((left - 24, y), f"{tick:.2f}", font=pil_font(27), fill=PIL_COLORS["gray"], anchor="rm")
    group_w = (right - left) / len(hybrid_rows)
    for idx, row in enumerate(hybrid_rows, start=1):
        center = left + (idx - 0.5) * group_w
        y_qr = bottom - int((float(row["delta_QR"]) - minimum) / (maximum - minimum) * (bottom - top))
        y_qc = bottom - int((float(row["delta_QC"]) - minimum) / (maximum - minimum) * (bottom - top))
        draw.ellipse((center - 14, y_qr - 14, center + 14, y_qr + 14), fill=PIL_COLORS["blue"])
        draw.polygon([(center, y_qc - 17), (center - 16, y_qc + 13), (center + 16, y_qc + 13)], fill=PIL_COLORS["orange"])
        draw.text((center, bottom + 20), str(idx), font=pil_font(23), fill=PIL_COLORS["dark"], anchor="ma")
    draw.text((1200, 1115), "任务序号", font=pil_font(29), fill=PIL_COLORS["dark"], anchor="mm")
    draw.ellipse((790, 1175, 830, 1215), fill=PIL_COLORS["blue"])
    draw.text((850, 1195), "C+R 与 C+Q 的差值", font=pil_font(27), fill=PIL_COLORS["dark"], anchor="lm")
    draw.polygon([(1260, 1175), (1240, 1215), (1280, 1215)], fill=PIL_COLORS["orange"])
    draw.text((1300, 1195), "纯经典与 C+Q 的差值", font=pil_font(27), fill=PIL_COLORS["dark"], anchor="lm")
    centered_text(draw, (1200, 1300), "24个任务的两类路线距离增量均为0。", pil_font(30, True), PIL_COLORS["red"])
    return save_chart(image, "图5_混合路线距离增量.png")


def set_run_font(run, size: float | None = None, bold: bool | None = None, color: str | None = None, italic: bool | None = None) -> None:
    run.font.name = "Microsoft YaHei"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Microsoft YaHei")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Microsoft YaHei")
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if italic is not None:
        run.italic = italic


def set_style_font(style, size: float, color: str = DARK, bold: bool = False) -> None:
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = bold
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Microsoft YaHei")
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Microsoft YaHei")
    style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: Sequence[int], indent_dxa: int = 120) -> None:
    if sum(widths_dxa) != 9360:
        raise ValueError(f"table widths must total 9360 DXA, got {sum(widths_dxa)}")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    grid_cols = table._tbl.tblGrid.findall(qn("w:gridCol"))
    for grid_col, width in zip(grid_cols, widths_dxa):
        grid_col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths_dxa):
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_text(cell, text: str, *, size: float = 9, bold: bool = False, color: str = DARK, align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.08
    run = p.add_run(str(text))
    set_run_font(run, size=size, bold=bold, color=color)


def shade_paragraph(paragraph, fill: str, border: str | None = None) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    if border:
        p_bdr = p_pr.find(qn("w:pBdr"))
        if p_bdr is None:
            p_bdr = OxmlElement("w:pBdr")
            p_pr.append(p_bdr)
        for edge in ("top", "left", "bottom", "right"):
            node = OxmlElement(f"w:{edge}")
            node.set(qn("w:val"), "single")
            node.set(qn("w:sz"), "8")
            node.set(qn("w:space"), "6")
            node.set(qn("w:color"), border)
            p_bdr.append(node)


def add_paragraph(doc: Document, text: str = "", *, bold: bool = False, color: str = DARK, size: float = 11, align=WD_ALIGN_PARAGRAPH.LEFT, before: float = 0, after: float = 6, line_spacing: float = 1.10, keep_with_next: bool = False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line_spacing
    p.paragraph_format.keep_with_next = keep_with_next
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size, bold=bold, color=color)
    return p


def add_labeled_paragraph(doc: Document, label: str, value: str, *, after: float = 3):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    r1 = p.add_run(f"{label}：")
    set_run_font(r1, size=10.5, bold=True, color=NAVY)
    r2 = p.add_run(value)
    set_run_font(r2, size=10.5, color=DARK)
    return p


def add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    set_run_font(run, bold=True, color=BLUE if level < 3 else NAVY)
    return p


def add_caption(doc: Document, text: str):
    p = add_paragraph(doc, text, size=9, color=MID_GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, before=3, after=8, keep_with_next=False)
    p.runs[0].italic = True
    return p


def add_picture(doc: Document, path: Path, caption: str, width: float = 6.15) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_together = True
    run = p.add_run()
    inline_shape = run.add_picture(str(path), width=Inches(width))
    inline_shape._inline.docPr.set("descr", caption)
    inline_shape._inline.docPr.set("title", caption.split("  ", 1)[0])
    add_caption(doc, caption)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, size=9, color=MID_GRAY)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    tail = paragraph.add_run(" 页")
    set_run_font(tail, size=9, color=MID_GRAY)


def add_metric_strip(doc: Document) -> None:
    metrics = [
        ("24/24", "真机任务完成"),
        ("24,576", "有效shots"),
        ("24/24", "命中最优BQM能量"),
        ("0/24", "严格能量胜出"),
    ]
    table = doc.add_table(rows=1, cols=4)
    set_table_geometry(table, [2340, 2340, 2340, 2340])
    repeat_table_header(table.rows[0])
    for cell, (value, label) in zip(table.rows[0].cells, metrics):
        set_cell_shading(cell, LIGHT_BLUE)
        cell.text = ""
        p1 = cell.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p1.paragraph_format.space_after = Pt(2)
        r1 = p1.add_run(value)
        set_run_font(r1, size=18, bold=True, color=NAVY)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(0)
        r2 = p2.add_run(label)
        set_run_font(r2, size=8.5, color=MID_GRAY)


def add_key_results_table(doc: Document) -> None:
    rows = [
        ("任务与shots", "24个任务；24,576/24,576 shots", "完整，无失败或不可评价任务"),
        ("真机QHR", "均值0.161011；95% CI [0.127035, 0.197795]", "反映低能量可行区域的实测命中频率"),
        ("随机参考", "均值0.226562；95% CI [0.186198, 0.270833]", "24个任务中真机超过随机参考为0次"),
        ("量子-随机", "均值-0.065552；95% CI [-0.079427, -0.051514]", "本协议下未观察到正向质量增益"),
        ("达到经典阈值", "平均比率0.050252", "所有任务均有样本达到或等于经典阈值"),
        ("最优BQM能量", "best_gap=0，24/24任务", "每个任务至少命中一次全局最优BQM能量"),
        ("严格经典改进", "0/24任务", "同预算经典方法已达到相同最优能量"),
        ("最终路线改善", "C+Q相对C+R为0/24任务", "量子候选没有改变最终路线距离"),
    ]
    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [1900, 3100, 4360])
    headers = ["指标", "正式结果", "解释"]
    for cell, text in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, NAVY)
        set_cell_text(cell, text, size=9.5, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
    repeat_table_header(table.rows[0])
    for metric, result, meaning in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], metric, size=8.8, bold=True, color=NAVY)
        set_cell_text(cells[1], result, size=8.5)
        set_cell_text(cells[2], meaning, size=8.5)
    set_table_geometry(table, [1900, 3100, 4360])


def add_instance_table(doc: Document) -> None:
    rows = [
        ("seed2026_c4_v2_medium", "4", "中等", "6", "Baihua / Dongling / Shenglian", "每后端2次"),
        ("seed2027_c4_v2_tight", "4", "紧张", "6", "Baihua / Dongling / Shenglian", "每后端2次"),
        ("seed2026_c6_v2_medium", "6", "中等", "9", "Baihua / Dongling / Shenglian", "每后端2次"),
        ("seed2027_c6_v2_tight", "6", "紧张", "8", "Baihua / Dongling / Shenglian", "每后端2次"),
    ]
    widths = [2700, 700, 900, 700, 3100, 1260]
    table = doc.add_table(rows=1, cols=6)
    set_table_geometry(table, widths)
    for cell, text in zip(table.rows[0].cells, ["实例ID", "客户数", "容量压力", "容量", "固定后端", "重复"]):
        set_cell_shading(cell, LIGHT_GRAY)
        set_cell_text(cell, text, size=8.5, bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER)
    repeat_table_header(table.rows[0])
    for row in rows:
        cells = table.add_row().cells
        for idx, (cell, text) in enumerate(zip(cells, row)):
            set_cell_text(cell, text, size=8.0, align=WD_ALIGN_PARAGRAPH.CENTER if idx != 0 else WD_ALIGN_PARAGRAPH.LEFT)
    set_table_geometry(table, widths)


def add_backend_table(doc: Document, stats: dict) -> None:
    table = doc.add_table(rows=1, cols=5)
    widths = [1700, 1000, 2200, 2200, 2260]
    set_table_geometry(table, widths)
    headers = ["后端", "任务数", "真机QHR均值", "随机参考均值", "量子-随机均值"]
    for cell, text in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, LIGHT_GRAY)
        set_cell_text(cell, text, size=8.5, bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER)
    repeat_table_header(table.rows[0])
    for backend in ("Baihua", "Dongling", "Shenglian"):
        entry = stats["by_backend"][backend]
        row = [
            backend,
            "8",
            f"{entry['quality_hit_rate']['mean']:.6f}",
            f"{entry['random_quality_hit_rate']['mean']:.6f}",
            f"{entry['quality_hit_gain']['mean']:.6f}",
        ]
        cells = table.add_row().cells
        for cell, text in zip(cells, row):
            set_cell_text(cell, text, size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_table_geometry(table, widths)


def short_instance(instance_id: str) -> str:
    mapping = {
        "seed2026_c4_v2_medium": "2026-C4-中等",
        "seed2027_c4_v2_tight": "2027-C4-紧张",
        "seed2026_c6_v2_medium": "2026-C6-中等",
        "seed2027_c6_v2_tight": "2027-C6-紧张",
    }
    return mapping.get(instance_id, instance_id)


def add_task_appendix(doc: Document, rows: list[dict[str, str]]) -> None:
    widths = [500, 1900, 1900, 1080, 620, 1000, 1000, 1360]
    table = doc.add_table(rows=1, cols=8)
    set_table_geometry(table, widths)
    headers = ["#", "任务ID", "实例", "后端", "轮次", "QHR", "随机参考", "经典阈值达到率"]
    for cell, text in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, NAVY)
        set_cell_text(cell, text, size=7.5, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
    repeat_table_header(table.rows[0])
    for idx, row in enumerate(rows, start=1):
        values = [
            str(idx),
            row["task_id"],
            short_instance(row["instance_id"]),
            row["backend_actual"],
            row["repeat_index"],
            f"{float(row['quality_hit_rate']):.6f}",
            f"{float(row['random_quality_hit_rate']):.6f}",
            f"{float(row['classical_reach_feasible_rate']):.6f}",
        ]
        cells = table.add_row().cells
        if idx % 2 == 0:
            for cell in cells:
                set_cell_shading(cell, "F8FAFC")
        for col, (cell, text) in enumerate(zip(cells, values)):
            set_cell_text(cell, text, size=7.2, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_table_geometry(table, widths)


def build_document(
    task_rows: list[dict[str, str]],
    hybrid_rows: list[dict[str, str]],
    stats: dict,
    charts: Sequence[Path],
) -> None:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    set_style_font(styles["Normal"], 11, DARK, False)
    styles["Normal"].paragraph_format.space_before = Pt(0)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.10
    set_style_font(styles["Title"], 24, NAVY, True)
    set_style_font(styles["Subtitle"], 13, MID_GRAY, False)
    set_style_font(styles["Heading 1"], 16, BLUE, True)
    styles["Heading 1"].paragraph_format.space_before = Pt(16)
    styles["Heading 1"].paragraph_format.space_after = Pt(8)
    styles["Heading 1"].paragraph_format.keep_with_next = True
    set_style_font(styles["Heading 2"], 13, BLUE, True)
    styles["Heading 2"].paragraph_format.space_before = Pt(12)
    styles["Heading 2"].paragraph_format.space_after = Pt(6)
    styles["Heading 2"].paragraph_format.keep_with_next = True
    set_style_font(styles["Heading 3"], 12, NAVY, True)
    styles["Heading 3"].paragraph_format.space_before = Pt(8)
    styles["Heading 3"].paragraph_format.space_after = Pt(4)
    styles["Heading 3"].paragraph_format.keep_with_next = True

    header_p = section.header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_p.paragraph_format.space_after = Pt(0)
    header_run = header_p.add_run("Quantum Route Forge  |  正式实验最终报告")
    set_run_font(header_run, size=8.5, bold=True, color=MID_GRAY)
    footer_p = section.footer.paragraphs[0]
    add_page_number(footer_p)

    add_paragraph(doc, "QUANTUM ROUTE FORGE / 正式硬件实验", size=10, bold=True, color=BLUE, after=8)
    title_p = doc.add_paragraph(style="Title")
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(6)
    title_run = title_p.add_run("Quantum Route Forge\n正式实验最终报告")
    set_run_font(title_run, size=24, bold=True, color=NAVY)
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.paragraph_format.space_after = Pt(16)
    subtitle_run = subtitle.add_run("24任务真实量子硬件矩阵、BQM候选质量与混合路线贡献分析")
    set_run_font(subtitle_run, size=13, color=MID_GRAY)
    add_labeled_paragraph(doc, "实验日期", "2026-08-02")
    add_labeled_paragraph(doc, "冻结协议", "formal-matrix-v2 / qrf-formal-protocol-v2")
    add_labeled_paragraph(doc, "执行版本", "9082b74 / qrf-preformal-execution-v2")
    add_labeled_paragraph(doc, "正式结果目录", "results/experiments/qrf_formal_hardware_matrix_v2")
    add_labeled_paragraph(doc, "发布状态", "Pull Request #1 已合并至 main；发布标签 qrf-final-experiment-v2.1", after=12)
    add_metric_strip(doc)
    add_paragraph(doc, "", after=6)
    callout = add_paragraph(
        doc,
        "核心结论  工程执行完整成功；真机在24/24任务中均命中过全局最优BQM能量，但没有严格击败同预算经典阈值，也没有改善最终路线距离。",
        bold=True,
        color=NAVY,
        size=11.5,
        before=8,
        after=8,
        line_spacing=1.25,
    )
    shade_paragraph(callout, "EAF2F8", BLUE)
    note = add_paragraph(
        doc,
        "统计说明：不把shots当作独立实验重复；置信区间以硬件任务为统计单位。所有不利结果均被保留。",
        size=9.5,
        color=MID_GRAY,
        before=8,
        after=0,
    )
    note.runs[0].italic = True

    doc.add_page_break()
    add_heading(doc, "一、执行摘要", 1)
    add_paragraph(
        doc,
        "本阶段按照预声明的4个实例、3个固定量子后端和2次重复执行正式硬件矩阵。24个任务全部完成，每个任务收到1024个有效shots，请求后端与实际后端完全一致，结果存储严格校验为complete=true、valid=true，错误和警告均为0。",
    )
    add_key_results_table(doc)
    add_heading(doc, "二、实验设计与可复现性", 1)
    add_paragraph(
        doc,
        "主实验采用两辆车和100%客户到量子比特覆盖。QAOA风格电路参数固定为gamma=1.1、beta=0.8；每个实例的客户顺序、逻辑QASM、经典阈值和配置哈希均在读取真机结果前冻结。每次命令最多提交一个新硬件任务，提交回执在轮询前持久化。",
    )
    add_instance_table(doc)
    add_paragraph(
        doc,
        "完整证据包括任务ID、计数分布、脱敏原始响应、逻辑QASM、客户顺序、冻结哈希、依赖快照、时间戳、队列/轮询字段和编译选项。正式分析只接纳source=hardware且状态为COMPLETED的数据。",
        after=0,
    )

    add_heading(doc, "三、量子候选质量结果", 1)
    add_paragraph(
        doc,
        "真机质量命中率QHR衡量采样结果进入预声明低能量可行区域的频率。24任务QHR均值为0.161011，冻结均匀随机参考均值为0.226562；配对差值均值为-0.065552，95%任务级Bootstrap置信区间为[-0.079427, -0.051514]。24个任务中没有任何任务的真机QHR超过随机参考。",
    )
    add_picture(doc, charts[0], "图1  逐任务真机质量命中率与冻结随机参考。")

    doc.add_page_break()
    add_heading(doc, "四、后端分层结果", 1)
    add_paragraph(
        doc,
        "三个后端均完成8个正式任务。后端分层仅作描述，不作为芯片优劣排名，因为本实验并未针对后端差异进行统计功效设计。",
    )
    add_backend_table(doc, stats)
    add_picture(doc, charts[1], "图2  各后端真机QHR与随机参考的描述性比较。")

    doc.add_page_break()
    add_heading(doc, "五、真机BQM能量是否胜出", 1)
    add_paragraph(
        doc,
        "需要区分“命中最优能量”和“严格击败经典方法”。best_gap定义为真机测得候选能量与该实例精确全局最优BQM能量之差。24个任务的best_gap均为0，因此每个任务至少有一个真机比特串命中全局最优BQM能量。",
    )
    add_paragraph(
        doc,
        "但严格改进率在24个任务中全部为0。结合best_gap=0，这说明同预算可行经典候选也已经达到相同的全局最优BQM能量；真机只能与经典最优打平，无法得到更低能量。在C+Q内部选择中，量子候选有11/24次成为最终来源，但这些属于相同能量或后续来源/排名规则下的选择，不能表述为严格能量胜出。",
    )
    add_picture(doc, charts[2], "图3  最优BQM能量命中、严格能量胜出、量子来源选择和路线改善的任务计数。")

    doc.add_page_break()
    add_heading(doc, "六、经典阈值达到情况", 1)
    add_paragraph(
        doc,
        "真机样本在每个任务中都以较低正比率达到或等于同预算可行经典阈值，整体平均达到率为0.050252；但没有样本严格低于该阈值。达到阈值代表候选质量与经典最好值相当，不代表量子方法击败经典方法。",
    )
    add_picture(doc, charts[3], "图4  四个冻结实例的经典阈值达到率和严格改进率。")

    doc.add_page_break()
    add_heading(doc, "七、公平混合路线贡献", 1)
    add_paragraph(
        doc,
        "公平比较中，C、C+R和C+Q使用相同总候选预算；C+R与C+Q共享完全相同的经典子集，仅将另外一半候选替换为随机候选或真机量子候选。所有候选经过相同的BQM评价、容量修复、路线构造、2-opt和最终选择。",
    )
    add_paragraph(
        doc,
        "24个任务的D_C+R-D_C+Q和D_C-D_C+Q均为0。换言之，真机量子候选没有在任何任务中缩短最终路线。量子来源在11个任务中被选中，说明其能够提供与最佳候选相同等级的种子，但没有产生可报告的路线距离增量。",
    )
    add_picture(doc, charts[4], "图5  24个任务的C+Q路线距离增量；所有点均位于0。")

    doc.add_page_break()
    add_heading(doc, "八、最终结论与声明边界", 1)
    add_heading(doc, "8.1 可以报告的结论", 2)
    add_paragraph(doc, "本次正式24任务真实硬件矩阵完整执行并通过严格验证。")
    add_paragraph(doc, "真机在所有任务中均命中过全局最优BQM能量，但未严格优于同预算经典阈值。")
    add_paragraph(doc, "真机QHR在24/24任务中低于冻结随机参考，未显示稳定的候选质量增益。")
    add_paragraph(doc, "C+Q没有在任何任务中改善最终路线距离。")
    add_heading(doc, "8.2 不应报告的结论", 2)
    add_paragraph(doc, "不能声称量子优势、速度优势、后端优劣排名或纯量子求解完整车辆路径问题。")
    add_paragraph(doc, "不能把能量打平、量子来源平局胜出或达到经典阈值改写为“量子击败经典”。")
    add_heading(doc, "8.3 适用范围", 2)
    add_paragraph(
        doc,
        "结论仅适用于本次冻结的4个小规模实例、两车辆设定、固定QAOA风格参数、三种被测后端、1024 shots和当前经典对照预算。结果不能外推为所有量子路线优化方法均无效。",
    )
    add_heading(doc, "九、审计与发布记录", 1)
    add_labeled_paragraph(doc, "完整性报告", "results/experiments/qrf_formal_hardware_matrix_v2/integrity_report.json")
    add_labeled_paragraph(doc, "统计汇总", "results/experiments/qrf_formal_hardware_matrix_v2/statistics_summary.json")
    add_labeled_paragraph(doc, "混合结果", "results/experiments/qrf_formal_hardware_matrix_v2/hybrid_summary.csv")
    add_labeled_paragraph(doc, "正式配置", "experiments/configs/formal_hardware_matrix_v2.json")
    add_labeled_paragraph(doc, "最终版本", "qrf-final-experiment-v2.1；PR #1已合并至main")
    add_paragraph(
        doc,
        "凭证扫描覆盖529个文本证据或源文件，命中数为0；415个正式证据与分析文件已在本地设置为只读。回归验证为44项测试通过，远端GitHub CI全部通过。",
        after=0,
    )

    doc.add_page_break()
    add_heading(doc, "附录A  24个正式硬件任务明细", 1)
    add_paragraph(
        doc,
        "表中QHR为真机质量命中率；随机参考为冻结均匀随机命中率；经典阈值达到率允许与经典阈值相等。24个任务的严格改进率均为0。",
        size=9.5,
        color=MID_GRAY,
        after=6,
    )
    add_task_appendix(doc, task_rows)

    core = doc.core_properties
    core.title = "Quantum Route Forge 正式实验最终报告（中文版）"
    core.subject = "24任务真实量子硬件矩阵、BQM候选质量与混合路线贡献分析"
    core.author = "Quantum Route Forge Project"
    core.keywords = "量子硬件, BQM, 路线优化, 正式实验, 可复现性"
    core.comments = "由冻结正式实验数据生成；不包含凭证。"

    DOCX_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(DOCX_PATH)


def main() -> None:
    if not FONT_REGULAR.exists() or not FONT_BOLD.exists():
        raise FileNotFoundError("Microsoft YaHei fonts are required for Chinese output")
    task_rows = load_csv(RESULT_ROOT / "candidate_quality_summary.csv")
    hybrid_rows = load_csv(RESULT_ROOT / "hybrid_summary.csv")
    stats = load_json(RESULT_ROOT / "statistics_summary.json")
    integrity = load_json(RESULT_ROOT / "integrity_report.json")
    if len(task_rows) != 24 or len(hybrid_rows) != 24:
        raise ValueError("expected exactly 24 formal task rows")
    if not integrity.get("complete") or not integrity.get("valid"):
        raise ValueError("formal result store is not complete and valid")
    if any(row["status"].lower() != "completed" for row in task_rows):
        raise ValueError("all formal task rows must be completed")

    charts = [
        chart_task_quality(task_rows),
        chart_backend_summary(stats),
        chart_bqm_outcomes(task_rows, hybrid_rows),
        chart_classical_threshold(stats),
        chart_hybrid_delta(hybrid_rows),
    ]
    build_document(task_rows, hybrid_rows, stats, charts)
    print(json.dumps({
        "docx": str(DOCX_PATH),
        "charts": [str(path) for path in charts],
        "tasks": len(task_rows),
        "integrity_valid": integrity["valid"],
    }, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
